import io
from datetime import datetime

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


_RED = RGBColor(0xC0, 0x00, 0x00)
_ORANGE = RGBColor(0xE6, 0x6C, 0x00)
_DARK = RGBColor(0x1A, 0x1A, 0x1A)


def export_playbook_to_docx(playbook, org_name: str = "Your Organization") -> bytes:
    """
    Render an OrgPlaybook as a .docx file.
    playbook: OrgPlaybook SQLAlchemy model (playbook.sections is list[dict]).
    Returns raw bytes of the .docx file.
    """
    doc = Document()

    # ── Title page ────────────────────────────────────────────────────────────
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run(f"{org_name}\nContract Standards Playbook")
    run.bold = True
    run.font.size = Pt(24)

    doc.add_paragraph()  # spacer

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(
        f"Version {playbook.version}  ·  "
        f"Generated {_fmt_date(playbook.generated_at)}  ·  "
        f"{playbook.doc_count} source documents"
    ).font.size = Pt(11)

    conf = doc.add_paragraph()
    conf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = conf.add_run("CONFIDENTIAL — INTERNAL USE ONLY")
    run.bold = True
    run.font.color.rgb = _RED

    doc.add_page_break()

    # ── Table of contents (manual) ────────────────────────────────────────────
    doc.add_heading("Table of Contents", level=1)
    sections = playbook.sections or []
    for i, section in enumerate(sections, 1):
        doc.add_paragraph(f"{i}. {section.get('title', section.get('clause_type', ''))}")
    doc.add_page_break()

    # ── Sections ──────────────────────────────────────────────────────────────
    for section in sections:
        _write_section(doc, section)
        doc.add_page_break()

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _write_section(doc: Document, section: dict) -> None:
    doc.add_heading(section.get("title", section.get("clause_type", "Section")), level=1)

    # Non-Negotiables
    non_neg = section.get("non_negotiables", [])
    if non_neg:
        doc.add_heading("Non-Negotiables", level=2)
        for item in non_neg:
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(item)
            run.font.color.rgb = _RED

    # Standard Positions
    positions = section.get("standard_positions", [])
    if positions:
        doc.add_heading("Standard Positions", level=2)
        table = doc.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        for cell, label in zip(hdr, ["Description", "Acceptable Range", "Rationale"]):
            cell.text = label
            cell.paragraphs[0].runs[0].bold = True

        for pos in positions:
            row = table.add_row().cells
            row[0].text = pos.get("description", "")
            row[1].text = pos.get("acceptable_range", "")
            row[2].text = pos.get("rationale", "")

    # Red Flags
    flags = section.get("red_flags", [])
    if flags:
        doc.add_heading("Red Flags", level=2)
        for flag in flags:
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(flag)
            run.font.color.rgb = _ORANGE

    # Industry Baseline
    baseline = section.get("industry_baseline", "")
    if baseline:
        doc.add_heading("Industry Baseline", level=2)
        doc.add_paragraph(baseline)


def _fmt_date(dt) -> str:
    if isinstance(dt, datetime):
        return dt.strftime("%B %d, %Y")
    if isinstance(dt, str):
        try:
            return datetime.fromisoformat(dt).strftime("%B %d, %Y")
        except ValueError:
            return dt
    return str(dt)
